##fthsn90##

from tkinter import *
import tkinter
from tkinter import messagebox
import random
from PIL import ImageTk,Image
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from docx import Document

doc = Document()
root = Tk()
root.geometry("800x320")
root.title("Dilekçe App")
root.resizable(width=FALSE, height=FALSE)
root.iconphoto(False, ImageTk.PhotoImage(Image.open("ikon.png")))
root.config(bg="#557c7f")
lst = ["Adı Soyadı","Başlık 1","Başlık 2","Tarih","Telefon","Adres","Ek"]
frmust = Frame(root,width=10,height=10,bg="#557c7f")
frmust.pack()
frmust.config(bd=2)
lbl = Label(frmust,text="DİLEKÇE YAZMA UYGULAMASI",font=("bold 10"),bg="#557c7f",fg="#ededed")
lbl.grid(row=0,column=1)
frmorta = Frame(root,bg="#557c7f")
frmorta.pack(fill=X)
menubar = Menu(root)
root.config(menu=menubar)
subMenu = Menu(menubar, tearoff=0)

def olusturbro():
    adsay = len(list(ad.get()))

    baslik = doc.add_paragraph(b1.get())
    paragraph_format1 = baslik.paragraph_format
    paragraph_format1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ikincibaslik = doc.add_paragraph(b2.get())
    paragraph_format2 = ikincibaslik.paragraph_format
    paragraph_format2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format2.left_indent = Inches(2)

    metin = doc.add_paragraph("             ")
    metin.space_before = Pt(10)
    metin.add_run("{}".format(mtn.get(1.0,END)))

    tarih = doc.add_paragraph(trh.get())
    paragraph_format4 = tarih.paragraph_format
    paragraph_format4.space_before = Pt(50)
    imza = doc.add_paragraph("İmza")
    adsoyad = doc.add_paragraph(ad.get())

    tarih = tarih.paragraph_format
    imza = imza.paragraph_format
    adsoyad = adsoyad.paragraph_format
    tarih.left_indent = Inches(5)
    # Ad soyadın uzunluğuna göre girintiyi ayarlama olduğu kadar.
    if adsay >= 17:
        adsoyad.left_indent = Inches(4.6)
    elif adsay == 9:
        adsoyad.left_indent = Inches(5.1)
    elif adsay <9:
        adsoyad.left_indent = Inches(5.2)
    elif adsay ==14:
        adsoyad.left_indent = Inches(4.8)
    elif adsay ==12:
        adsoyad.left_indent = Inches(4.9)
    else:
        adsoyad.left_indent = Inches(4.7) 
    imza.left_indent = Inches(5.3)	
    if adrs.get() and tlf.get() and ek.get():
        adres = doc.add_paragraph("Adres: {}".format(adrs.get()))
        paragraph_format3 = adres.paragraph_format
        paragraph_format3.space_before = Pt(100)
        Telefon = doc.add_paragraph("Telefon: {}".format(tlf.get()))
        Ek = doc.add_paragraph("Ek : {}".format(ek.get()))
    if ek.get() and not adrs.get() and not tlf.get():
        Ek = doc.add_paragraph("Ek : {}".format(ek.get()))
        paragraph_format5 = Ek.paragraph_format
        paragraph_format5.space_before = Pt(100)

    messagebox.showinfo("Bilgi","Dilekçe başarıyla oluşturuldu!")
    doc.save("dilekce.docx")

def Hakkında():
    messagebox.showinfo("Bilgi","İşbu program @fthsn tarafından kodlanmıştır. Github: fthsn90")
x= 0
for i in range(7):
    lbl = Label(frmorta,text=f"{lst[i]}",bg="#557c7f",fg="#ededed")
    lbl.grid(column=x,row=i,padx=5,pady=5)

ad = Entry(frmorta,bd=2,bg="#557c7f",fg="#ededed")
ad.grid(row=0,column=1,padx=8)
b1 = Entry(frmorta,bd=2,bg="#557c7f",fg="#ededed")
b1.grid(row=1,column=1)
b2 = Entry(frmorta,bd=2,bg="#557c7f",fg="#ededed")
b2.grid(row=2,column=1)
trh = Entry(frmorta,bd=2,bg="#557c7f",fg="#ededed")
trh.grid(row=3,column=1)
tlf = Entry(frmorta,bd=2,bg="#557c7f",fg="#ededed")
tlf.grid(row=4,column=1)
adrs = Entry(frmorta,bd=2,bg="#557c7f",fg="#ededed")
adrs.grid(row=5,column=1)
ek = Entry(frmorta,bd=2,bg="#557c7f",fg="#ededed")
ek.grid(row=6,column=1)
lblmetin =Label(frmorta,text="    METİN",bg="#557c7f",fg="#ededed")
lblmetin.grid(row=0,column=2)
mtn = Text(frmorta,width=72,height=14,bd=2,bg="#557c7f",fg="#ededed")
mtn.grid(row=1,column=2,rowspan=10)
btn = Button(frmorta,text = "OLUŞTUR",cursor="tcross",bd=5,bg="#557c7f",fg="#ededed",command=olusturbro)
btn.grid(row=8,column=1)
menubar.add_cascade(label="Bilgi", menu=subMenu)
subMenu.add_command(label="Hakkında", command=Hakkında)
root.mainloop()